import pdfplumber
import docx
from flask import Flask, request, jsonify, render_template
import pickle
import numpy as np

app = Flask(__name__)

# =========================
# Load ML model
# =========================
model = pickle.load(open("career_model.pkl", "rb"))
label_encoder = pickle.load(open("label_encoder.pkl", "rb"))

@app.route("/")
def home():
    return render_template("index.html")

# =========================
# ML FEATURES (MUST MATCH TRAINING)
# =========================
skills_list = [
    "python","java","sql","ml","excel",
    "communication","design","cloud",
    "marketing","management","aws",
    "docker","linux","sap","cyber",
    "react","seo"
]

# =========================
# ROLE → SKILL MAP
# =========================
role_skill_map = {
    # TECH
    "Data Scientist": ["python","ml","sql","statistics","pandas","numpy"],
    "Machine Learning Engineer": ["python","ml","deep learning","tensorflow","pytorch"],
    "AI Engineer": ["python","ai","nlp","opencv","ml"],
    "AI Researcher": ["python","deep learning","nlp","pytorch","tensorflow"],
    "Full Stack Developer": ["html","css","javascript","python","java","sql"],
    "Backend Developer": ["python","java","sql","api","flask","django"],
    "Frontend Developer": ["html","css","javascript","react","ui"],
    "UI/UX Designer": ["figma","design","ui","ux","wireframe"],
    "Cloud Engineer": ["aws","azure","cloud","linux","docker"],
    "Cloud Architect": ["aws","azure","cloud","devops","terraform","kubernetes"],
    "DevOps Engineer": ["docker","kubernetes","linux","ci/cd","cloud"],
    "Cyber Security Analyst": ["cyber","security","soc","siem","vulnerability","pentest"],
    "Network Engineer": ["network","router","switch","tcp","ip","ccna"],
    "SAP Consultant": ["sap","abap","hana","fico","mm","sd"],
    "Blockchain Developer": ["blockchain","solidity","smart contract","web3"],
    "Game Developer": ["unity","c#","game","3d","unreal"],

    # NON TECH
    "Business Analyst": ["excel","sql","communication","analysis","powerbi"],
    "Digital Marketer": ["seo","content","marketing","analytics","ads"],
    "HR Executive": ["recruitment","communication","interview","onboarding"],
    "Sales Executive": ["sales","communication","crm","negotiation"],
    "Operations Manager": ["management","operations","planning","coordination"],
    "Project Manager": ["management","planning","risk","stakeholder"],
    "Content Writer": ["writing","content","blog","seo"],
    "Customer Support": ["communication","support","client","service"],
    "Finance Analyst": ["finance","excel","accounting","budget"],
    "Supply Chain Analyst": ["logistics","inventory","planning","operations"],
    "Product Manager": ["roadmap","stakeholder","planning","market"]
}

# =========================
# DOMAIN DETECTION
# =========================
domain_map = {
    "Cyber Security": ["cyber","security","pentest","soc","siem","firewall"],
    "Cloud": ["aws","azure","gcp","cloud","ec2","s3","devops"],
    "SAP": ["sap","abap","hana","fico","mm","sd","basis"],
    "Networking": ["network","ccna","router","switch","tcp","ip"],
    "AI/ML": ["machine learning","deep learning","tensorflow","nlp","ml"],
    "Software Dev": ["developer","coding","python","java","api","git"],
    "Non-Tech": ["hr","sales","marketing","business","management","operations"]
}

# =========================
# TEXT EXTRACTION
# =========================
def extract_text(file):
    name = file.filename.lower()
    file.stream.seek(0)

    if name.endswith(".pdf"):
        text = ""
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
        return text.lower()

    elif name.endswith(".docx"):
        doc = docx.Document(file)
        return " ".join(p.text for p in doc.paragraphs).lower()

    return ""

# =========================
# STRUCTURE CHECK
# =========================
def get_counts_from_file(file, text):
    resume_keywords = [
        "experience","education","skills","projects","internship",
        "certification","objective","summary","profile","achievements",
        "responsibilities","company","degree","university","technical",
        "developer","engineer","analyst","manager","worked","role"
    ]

    word_hits = sum(1 for w in resume_keywords if w in text)
    word_count = len(text.split())

    file.stream.seek(0)
    if file.filename.lower().endswith(".docx"):
        doc = docx.Document(file)
        structure_count = len([p for p in doc.paragraphs if p.text.strip()])
    else:
        structure_count = len([ln for ln in text.splitlines() if ln.strip()])

    return word_hits, word_count, structure_count

# =========================
# PREDICT ROUTE
# =========================
@app.route("/predict", methods=["POST"])
def predict():
    if "resume" not in request.files:
        return jsonify({"error": "No file uploaded"})

    file = request.files["resume"]
    text = extract_text(file)

    if not text or len(text.strip()) < 50:
        return jsonify({"error": "Wrong file uploaded. Please upload a valid resume (PDF/DOCX)."})

    word_hits, word_count, structure_count = get_counts_from_file(file, text)

    if word_hits < 3 or word_count < 120 or structure_count < 6:
        return jsonify({"error": "Wrong file uploaded. Please upload a valid resume/CV."})

    # ---- ML FEATURES ----
    features = [1 if skill in text else 0 for skill in skills_list]
    X = np.array(features).reshape(1, -1)
    probs = model.predict_proba(X)[0]

    # ---- TOP 3 ----
    top3_idx = probs.argsort()[-3:][::-1]
    top3_roles = label_encoder.inverse_transform(top3_idx)
    top3_scores = [round(probs[i] * 100, 2) for i in top3_idx]
    top3 = [{"role": r, "score": s} for r, s in zip(top3_roles, top3_scores)]

    # ---- PROGRESS ----
    progress = []
    for role, req_skills in role_skill_map.items():
        have = sum(1 for s in req_skills if s in text)
        total = len(req_skills)
        percent = round((have / total) * 100, 2) if total else 0
        progress.append({"role": role, "percent": percent})

    # ---- BEST TARGET ----
    best_target = None
    missing_skills = []

    for i in probs.argsort()[::-1]:
        if probs[i] < 0.50:
            role = label_encoder.inverse_transform([i])[0]
            required = role_skill_map.get(role, [])
            missing_skills = [s for s in required if s not in text]
            best_target = role
            break

    # ---- DOMAIN ----
    detected_domains = []
    for domain, keys in domain_map.items():
        if any(k in text for k in keys):
            detected_domains.append(domain)

    if not detected_domains:
        detected_domains = ["General / Unknown"]

    return jsonify({
        "top3": top3,
        "best_target": best_target,
        "missing_skills": missing_skills[:6],
        "progress": progress,
        "detected_domains": detected_domains
    })

if __name__ == "__main__":
    app.run(debug=True)
